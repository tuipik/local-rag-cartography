"""Text extraction pipeline for cataloged documents.

This stage deliberately does not run OCR, embeddings, or LLM calls. It only
extracts available machine-readable text and stores it in SQLite.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from urllib.request import pathname2url

from local_rag.database import (
    add_database_argument,
    connect_database,
    resolve_database,
)


SUPPORTED_EXTENSIONS = {"pdf", "doc", "docx", "txt"}
DEFAULT_DOC_TIMEOUT_SECONDS = 60


@dataclass(frozen=True)
class DocumentRecord:
    id: int
    path: Path
    relative_path: str
    extension: str


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str
    extraction_method: str


@dataclass(frozen=True)
class ExtractionResult:
    document: DocumentRecord
    pages: list[ExtractedPage]

    @property
    def has_text(self) -> bool:
        return any(page.text.strip() for page in self.pages)

    @property
    def ocr_required(self) -> bool:
        return self.document.extension == "pdf" and not self.has_text

    @property
    def status(self) -> str:
        return "ocr_required" if self.ocr_required else "text_extracted"


class ExtractionError(RuntimeError):
    """Raised when a document cannot be processed by the current stage."""


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract text from cataloged documents into SQLite."
    )
    add_database_argument(parser)
    parser.add_argument(
        "--document-id",
        type=int,
        help="extract text for a single document id",
    )
    parser.add_argument(
        "--doc-timeout-seconds",
        type=int,
        default=DEFAULT_DOC_TIMEOUT_SECONDS,
        help=(
            "maximum LibreOffice conversion time for one .doc file "
            f"(default: {DEFAULT_DOC_TIMEOUT_SECONDS})"
        ),
    )
    return parser.parse_args()


def initialize_database(connection: sqlite3.Connection) -> None:
    connection.execute("PRAGMA foreign_keys = ON")
    connection.execute(
        """
        CREATE TABLE IF NOT EXISTS document_pages (
            id INTEGER PRIMARY KEY,
            document_id INTEGER NOT NULL,
            page_number INTEGER NOT NULL,
            text TEXT NOT NULL,
            extraction_method TEXT NOT NULL,
            text_length INTEGER NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(document_id) REFERENCES documents(id),
            UNIQUE(document_id, page_number)
        )
        """
    )
    connection.execute(
        """
        CREATE INDEX IF NOT EXISTS idx_document_pages_document_id
        ON document_pages(document_id)
        """
    )


def fetch_documents(
    connection: sqlite3.Connection, document_id: int | None = None
) -> list[DocumentRecord]:
    where_clause = """
        extension IN ('pdf', 'doc', 'docx', 'txt')
        AND COALESCE(is_supported, 1) = 1
    """
    params: list[object] = []
    if document_id is not None:
        where_clause += " AND id = ?"
        params.append(document_id)

    rows = connection.execute(
        f"""
        SELECT id, path, relative_path, extension
        FROM documents
        WHERE {where_clause}
        ORDER BY id
        """,
        params,
    ).fetchall()
    return [
        DocumentRecord(
            id=row["id"],
            path=Path(row["path"]),
            relative_path=row["relative_path"],
            extension=row["extension"].lower(),
        )
        for row in rows
    ]


def extract_pdf(path: Path) -> list[ExtractedPage]:
    try:
        import fitz
    except ImportError as error:
        raise ExtractionError("PyMuPDF is not installed") from error

    try:
        with fitz.open(path) as document:
            return [
                ExtractedPage(
                    page_number=index + 1,
                    text=page.get_text("text") or "",
                    extraction_method="pymupdf",
                )
                for index, page in enumerate(document)
            ]
    except Exception as error:
        raise ExtractionError(str(error)) from error


def extract_docx(path: Path) -> list[ExtractedPage]:
    try:
        from docx import Document
    except ImportError as error:
        raise ExtractionError("python-docx is not installed") from error

    try:
        document = Document(path)
    except Exception as error:
        raise ExtractionError(str(error)) from error

    blocks: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            blocks.append("\t".join(cell for cell in cells if cell))

    text = "\n".join(block for block in blocks if block.strip())
    return [ExtractedPage(1, text, "python-docx")]


def read_text_file(path: Path) -> str:
    encodings = ("utf-8-sig", "utf-8", "cp1251", "cp1252")
    last_error: UnicodeDecodeError | None = None
    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as error:
            last_error = error
    if last_error is not None:
        raise ExtractionError(str(last_error)) from last_error
    raise ExtractionError("unable to read text file")


def extract_txt(path: Path) -> list[ExtractedPage]:
    return [ExtractedPage(1, read_text_file(path), "plain-text")]


def libreoffice_binary() -> str | None:
    return shutil.which("soffice") or shutil.which("libreoffice")


def extract_doc(
    path: Path, timeout_seconds: int = DEFAULT_DOC_TIMEOUT_SECONDS
) -> list[ExtractedPage]:
    binary = libreoffice_binary()
    if binary is None:
        raise ExtractionError("LibreOffice is not installed")

    with tempfile.TemporaryDirectory(prefix="doc-extract-") as temp_directory:
        output_dir = Path(temp_directory)
        profile_dir = output_dir / "libreoffice-profile"
        command = [
            binary,
            "--headless",
            f"-env:UserInstallation=file://{pathname2url(str(profile_dir))}",
            "--convert-to",
            "txt:Text",
            "--outdir",
            str(output_dir),
            str(path),
        ]
        try:
            result = subprocess.run(
                command,
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as error:
            raise ExtractionError(
                f"LibreOffice conversion timed out after {timeout_seconds}s"
            ) from error
        if result.returncode != 0:
            message = "\n".join(
                part
                for part in (result.stderr.strip(), result.stdout.strip())
                if part
            )
            raise ExtractionError(message or "LibreOffice conversion failed")

        converted = output_dir / f"{path.stem}.txt"
        if not converted.exists():
            candidates = list(output_dir.glob("*.txt"))
            if not candidates:
                raise ExtractionError("LibreOffice did not create a text file")
            converted = candidates[0]

        return [ExtractedPage(1, read_text_file(converted), "libreoffice-txt")]


def extract_document(
    document: DocumentRecord,
    doc_timeout_seconds: int = DEFAULT_DOC_TIMEOUT_SECONDS,
) -> ExtractionResult:
    if not document.path.exists():
        raise ExtractionError("file does not exist")
    if document.extension == "pdf":
        pages = extract_pdf(document.path)
    elif document.extension == "docx":
        pages = extract_docx(document.path)
    elif document.extension == "doc":
        pages = extract_doc(document.path, timeout_seconds=doc_timeout_seconds)
    elif document.extension == "txt":
        pages = extract_txt(document.path)
    else:
        raise ExtractionError(f"unsupported extension: {document.extension}")

    if not pages:
        pages = [ExtractedPage(1, "", "empty")]
    return ExtractionResult(document=document, pages=pages)


def save_result(connection: sqlite3.Connection, result: ExtractionResult) -> None:
    created_at = utc_now()
    connection.execute(
        "DELETE FROM document_pages WHERE document_id = ?",
        (result.document.id,),
    )
    connection.executemany(
        """
        INSERT INTO document_pages (
            document_id, page_number, text, extraction_method, text_length,
            created_at
        ) VALUES (?, ?, ?, ?, ?, ?)
        """,
        [
            (
                result.document.id,
                page.page_number,
                page.text,
                page.extraction_method,
                len(page.text),
                created_at,
            )
            for page in result.pages
        ],
    )
    connection.execute(
        """
        UPDATE documents
        SET pages_count = ?,
            has_text = ?,
            ocr_required = ?,
            scan_status = ?
        WHERE id = ?
        """,
        (
            len(result.pages),
            int(result.has_text),
            int(result.ocr_required),
            result.status,
            result.document.id,
        ),
    )


def save_error(
    connection: sqlite3.Connection, document: DocumentRecord, error: Exception
) -> None:
    connection.execute(
        "DELETE FROM document_pages WHERE document_id = ?",
        (document.id,),
    )
    connection.execute(
        """
        UPDATE documents
        SET pages_count = NULL,
            has_text = 0,
            ocr_required = CASE WHEN extension = 'pdf' THEN 1 ELSE 0 END,
            scan_status = ?
        WHERE id = ?
        """,
        (f"extraction_error: {error}", document.id),
    )


def run_extraction(
    database: Path,
    document_id: int | None = None,
    doc_timeout_seconds: int = DEFAULT_DOC_TIMEOUT_SECONDS,
) -> Counter[str]:
    stats: Counter[str] = Counter()
    with connect_database(database) as connection:
        initialize_database(connection)
        documents = fetch_documents(connection, document_id=document_id)
        stats["documents"] = len(documents)

        for document in documents:
            stats[f"{document.extension}_processed"] += 1
            try:
                result = extract_document(
                    document,
                    doc_timeout_seconds=doc_timeout_seconds,
                )
                save_result(connection, result)
                if not result.has_text:
                    stats["empty_text"] += 1
                if result.ocr_required:
                    stats["ocr_required"] += 1
            except Exception as error:
                save_error(connection, document, error)
                stats["errors"] += 1
                print(f"ERROR {document.relative_path}: {error}")

    return stats


def print_statistics(stats: Counter[str]) -> None:
    print("\nСтатистика витягу тексту")
    print(f"Documents processed: {stats['documents']}")
    print(f"PDF processed: {stats['pdf_processed']}")
    print(f"DOCX processed: {stats['docx_processed']}")
    print(f"DOC processed: {stats['doc_processed']}")
    print(f"TXT processed: {stats['txt_processed']}")
    print(f"empty_text: {stats['empty_text']}")
    print(f"ocr_required: {stats['ocr_required']}")
    print(f"errors: {stats['errors']}")


def main() -> int:
    args = parse_args()
    database = resolve_database(args.database)

    try:
        stats = run_extraction(
            database,
            document_id=args.document_id,
            doc_timeout_seconds=args.doc_timeout_seconds,
        )
    except FileNotFoundError as error:
        print(error)
        return 2
    print_statistics(stats)
    return 0 if stats["errors"] == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
